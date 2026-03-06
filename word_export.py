import subprocess
import json
import os
import shutil
import tempfile
from datetime import datetime

TEMP_DIR  = tempfile.gettempdir()
TEMP_JSON = os.path.join(TEMP_DIR, 'comparison_data.json')
TEMP_DOCX = os.path.join(TEMP_DIR, 'comparison_output.docx')


def export_to_word(result: str, query: str = ""):
    """
    Primary export — uses generate_word.js (Node.js) for the styled report.
    Falls back to export_to_word_basic() if Node.js is unavailable.

    `result` is expected to contain === SECTION 3: KEY SUMMARY === output
    from the new llm_engine.py, but also handles the old 4-section format.
    """
    data = {
        "query":        query,
        "result":       result,
        "generated_on": datetime.now().strftime("%d %B %Y, %I:%M %p"),
        "annexures":    []
    }

    try:
        with open(TEMP_JSON, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)

        script_dir = os.path.dirname(os.path.abspath(__file__))
        js_script  = os.path.join(script_dir, 'generate_word.js')

        if not os.path.exists(js_script):
            raise FileNotFoundError(f"generate_word.js not found at {js_script}")

        proc = subprocess.run(
            ['node', js_script],
            capture_output=True, text=True, timeout=60,
            cwd=script_dir
        )

        if proc.returncode == 0 and os.path.exists(TEMP_DOCX):
            output_path = os.path.join(script_dir, 'Income_Tax_Summary.docx')
            shutil.copy(TEMP_DOCX, output_path)
            return output_path
        else:
            raise Exception(proc.stderr or "Node.js returned non-zero exit code")

    except Exception as e:
        print(f"[TEJAS] Node.js export failed: {e} — falling back to python-docx")
        return export_to_word_basic(result, query)


# ── Sub-section definitions matching generate_word.js ────────────────────────
_SUB_SECTIONS = [
    "What this provision does",
    "Who it applies to",
    "The rules explained simply",
    "Key thresholds and rates at a glance",
    "What happens if you don't comply",
    "Worked example",
    "Note on 2025 Act structure",
]


def _extract_sec3(result: str) -> str:
    """
    Extract KEY SUMMARY content from result.
    Handles both new format (=== SECTION 3: KEY SUMMARY ===)
    and old format (=== SECTION 3 ===).
    Falls back to entire result if no marker found.
    """
    import re
    # New format
    m = re.search(r'===\s*SECTION\s*3[^=]*===([\s\S]*?)(?====|$)', result, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    # Old format fallback
    start = result.find('SECTION 3')
    if start != -1:
        end = result.find('SECTION 4', start)
        return result[start:end].strip() if end != -1 else result[start:].strip()
    # Last resort — return everything
    return result.strip()


def _split_sub_sections(text: str) -> dict:
    """
    Split KEY SUMMARY text into named sub-sections by **Bold heading** markers.
    Returns { heading: content_text }
    """
    import re
    result   = {}
    current  = '__preamble__'
    buffer   = []

    for line in text.split('\n'):
        m = re.match(r'^\*\*(.+?)\*\*\s*$', line.strip())
        if m:
            if buffer:
                result[current] = '\n'.join(buffer).strip()
            current = m.group(1).strip()
            buffer  = []
        else:
            buffer.append(line)

    if buffer:
        result[current] = '\n'.join(buffer).strip()

    return result


def export_to_word_basic(result: str, query: str = ""):
    """
    Pure-Python fallback using python-docx.
    Renders the KEY SUMMARY with sub-section headings.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading('TaxGenie — Income Tax Act 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading('Key Summary Report', level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Section / Topic: {query}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    doc.add_paragraph(f"Source: Income Tax Act 2025 (actual statutory text)")
    doc.add_paragraph("")

    # ── Extract and split content ──────────────────────────────────────────
    sec3        = _extract_sec3(result)
    sub_sections = _split_sub_sections(sec3)

    # Render each named sub-section, then any leftovers
    rendered_keys = set()

    for heading in _SUB_SECTIONS:
        # Partial match — tolerant of minor wording differences
        matched_key = next(
            (k for k in sub_sections if heading.lower()[:12] in k.lower()),
            None
        )
        content = sub_sections.get(matched_key, '') if matched_key else ''
        if not content:
            continue

        rendered_keys.add(matched_key)
        doc.add_heading(heading, level=1)

        for line in content.split('\n'):
            t = line.strip()
            if not t:
                continue

            # Bullet point
            if re.match(r'^[*\-]\s', t):
                clean = re.sub(r'\*\*', '', t[2:].strip())
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(clean).font.size = Pt(11)

            # Numbered list
            elif re.match(r'^\d+\.\s', t):
                clean = re.sub(r'\*\*', '', re.sub(r'^\d+\.\s*', '', t))
                p = doc.add_paragraph(style='List Number')
                p.add_run(clean).font.size = Pt(11)

            # **Label**: value
            elif re.match(r'^\*\*.+\*\*[:\s—–-]', t):
                m = re.match(r'^\*\*(.+?)\*\*[:\s—–-]+(.*)', t)
                if m:
                    p = doc.add_paragraph()
                    bold_run = p.add_run(m.group(1) + ': ')
                    bold_run.bold = True
                    bold_run.font.size = Pt(11)
                    p.add_run(re.sub(r'\*\*', '', m.group(2))).font.size = Pt(11)

            # **Bold heading** on its own line
            elif re.match(r'^\*\*[^*]+\*\*\s*$', t):
                h = doc.add_heading(re.sub(r'\*\*', '', t), level=3)

            # Normal paragraph — strip inline bold markers
            else:
                clean = re.sub(r'\*\*', '', t)
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(4)

        doc.add_paragraph("")

    # Any content not matched to a known sub-section (preamble etc.)
    for key, content in sub_sections.items():
        if key not in rendered_keys and key != '__preamble__' and content.strip():
            doc.add_heading(key, level=1)
            for line in content.split('\n'):
                t = line.strip()
                if t:
                    doc.add_paragraph(re.sub(r'\*\*', '', t))

    # ── Footer note ────────────────────────────────────────────────────────
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "This report is based on the actual text of the Income Tax Act 2025. "
        "Generated by TaxGenie"
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

    script_dir  = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, 'Income_Tax_Summary.docx')
    doc.save(output_path)
    return output_path